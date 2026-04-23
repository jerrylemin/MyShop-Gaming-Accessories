from __future__ import annotations

import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


INLINE_PATTERN = re.compile(r"(\*\*.*?\*\*|`.*?`)")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")
NUMBERED_PATTERN = re.compile(r"^\d+\.\s+(.*)$")


def parse_project_title(markdown_text: str) -> str:
    match = re.search(r"^\*\*(.+?)\*\*$", markdown_text, re.MULTILINE)
    if not match:
        raise ValueError("Could not extract project title from markdown.")
    return match.group(1).strip()


def parse_members(markdown_lines: list[str]) -> list[tuple[str, str]]:
    members: list[tuple[str, str]] = []
    for line in markdown_lines:
        if not line.startswith("|"):
            continue

        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if len(cells) != 3 or not cells[0].isdigit():
            continue

        members.append((cells[1], cells[2]))

    if not members:
        raise ValueError("Could not extract member table from markdown.")

    return members


def infer_academic_year(today: datetime) -> str:
    if today.month >= 8:
        return f"{today.year} - {today.year + 1}"
    return f"{today.year - 1} - {today.year}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)

    document.styles["Title"].font.size = Pt(16)
    document.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.styles["Heading 1"].font.size = Pt(12)
    document.styles["Heading 2"].font.size = Pt(12)
    document.styles["Heading 3"].font.size = Pt(12)

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(0)

    if "BodyTextProposal" not in document.styles:
        body_style = document.styles.add_style("BodyTextProposal", WD_STYLE_TYPE.PARAGRAPH)
        body_style.base_style = document.styles["Normal"]
        body_style.font.name = "Times New Roman"
        body_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        body_style.paragraph_format.space_after = Pt(6)


def add_inline_runs(paragraph, text: str) -> None:
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Times New Roman"
        else:
            run = paragraph.add_run(part)

        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_cover_page(document: Document, project_title: str, members: list[tuple[str, str]], academic_year: str) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ĐỀ XUẤT DỰ ÁN")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(16)

    info_lines = [
        f"Tên đề tài: {project_title}",
        "Môn học: Lập trình Windows",
        "Nhóm sinh viên:",
    ]

    for info_line in info_lines:
        paragraph = document.add_paragraph(style="BodyTextProposal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(paragraph, info_line)

    for student_id, student_name in members:
        paragraph = document.add_paragraph(style="BodyTextProposal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(paragraph, f"{student_id} - {student_name}")

    academic_year_paragraph = document.add_paragraph(style="BodyTextProposal")
    academic_year_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(academic_year_paragraph, f"Năm học: {academic_year}")

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"

    for row_index, row_values in enumerate(rows):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles["BodyTextProposal"]
            add_inline_runs(paragraph, value)
            if row_index == 0:
                set_cell_shading(cell, "D9E2F3")
                for run in paragraph.runs:
                    run.bold = True

    document.add_paragraph(style="BodyTextProposal")


def parse_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start_index
    while index < len(lines) and lines[index].strip().startswith("|"):
        current_line = lines[index].strip()
        if set(current_line.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set():
            index += 1
            continue

        cells = [cell.strip() for cell in current_line.strip("|").split("|")]
        rows.append(cells)
        index += 1

    return rows, index


def parse_paragraph(lines: list[str], start_index: int) -> tuple[str, int]:
    parts = [lines[start_index].strip()]
    index = start_index + 1
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            break
        if stripped.startswith("|") or stripped.startswith("- "):
            break
        if HEADING_PATTERN.match(stripped) or NUMBERED_PATTERN.match(stripped):
            break
        parts.append(stripped)
        index += 1

    return " ".join(parts), index


def render_markdown(document: Document, markdown_lines: list[str]) -> None:
    index = 0
    skipped_first_title = False
    while index < len(markdown_lines):
        raw_line = markdown_lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()
            if level == 1 and title_text == "ĐỀ XUẤT DỰ ÁN" and not skipped_first_title:
                skipped_first_title = True
                index += 1
                continue

            style_name = {
                1: "Heading 1",
                2: "Heading 1",
                3: "Heading 2",
            }.get(level, "Heading 3")
            paragraph = document.add_paragraph(style=style_name)
            add_inline_runs(paragraph, title_text)
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(markdown_lines, index)
            add_table(document, rows)
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:].strip())
            index += 1
            continue

        numbered_match = NUMBERED_PATTERN.match(stripped)
        if numbered_match:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, numbered_match.group(1).strip())
            index += 1
            continue

        paragraph_text, index = parse_paragraph(markdown_lines, index)
        paragraph = document.add_paragraph(style="BodyTextProposal")
        add_inline_runs(paragraph, paragraph_text)


def validate_output(docx_path: Path) -> None:
    document = Document(docx_path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if "# " in full_text or "## " in full_text or "```" in full_text:
        raise ValueError("Markdown heading or fence syntax remained in the DOCX output.")

    if not any(table.rows for table in document.tables):
        raise ValueError("Expected tables were not found in the DOCX output.")

    required_phrases = [
        "ĐỀ XUẤT DỰ ÁN",
        "Môn học: Lập trình Windows",
        "MyShop Gaming Accessories POS",
        "Lê Minh",
        "Nguyễn Vũ Bách",
    ]
    for phrase in required_phrases:
        if phrase not in full_text:
            raise ValueError(f"Expected phrase missing from DOCX output: {phrase}")


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python tools/generate_proposal_docx.py <input_md> <output_docx>")
        return 1

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    if not input_path.exists():
        print(f"Input markdown file not found: {input_path}")
        return 1

    markdown_text = input_path.read_text(encoding="utf-8")
    markdown_lines = markdown_text.splitlines()

    document = Document()
    configure_document(document)

    project_title = parse_project_title(markdown_text)
    members = parse_members(markdown_lines)
    academic_year = infer_academic_year(datetime.now())

    add_cover_page(document, project_title, members, academic_year)
    render_markdown(document, markdown_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    validate_output(output_path)
    print(f"Generated DOCX: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
